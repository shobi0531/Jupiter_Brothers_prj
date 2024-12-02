import sys
import os
import uuid
from pdf2docx import Converter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

def extract_tables_from_docx(docx_path):
    doc = Document(docx_path)
    tables = doc.tables

    valid_tables = []
    for table in tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        df = pd.DataFrame(table_data)
        if df.shape[0] > 1 and df.shape[1] > 1:
            valid_tables.append(df)
    return valid_tables

def set_first_row_as_header(table_df):
    new_header = table_df.iloc[0]
    table_df = table_df[1:]
    table_df.columns = new_header
    return table_df

def process_pdf_to_excel(input_pdf_path, output_dir):
    """Converts PDF to DOCX, processes tables, and saves to Excel"""
    unique_id = str(uuid.uuid4()) 
    docx_file = f"{output_dir}/output_{unique_id}.docx"

    cv = Converter(input_pdf_path)
    cv.convert(docx_file)
    cv.close()

    doc = Document(docx_file)
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    center_aligned_docx = f"{output_dir}/center_aligned_{unique_id}.docx"
    doc.save(center_aligned_docx)

    table_dfs = extract_tables_from_docx(center_aligned_docx)

    if not table_dfs:
        print("No tables found in the document.")
        return None

    processed_tables = []
    for i, table_df in enumerate(table_dfs):
        table_df = set_first_row_as_header(table_df)
        table_df.reset_index(drop=True, inplace=True)
        processed_tables.append(table_df)

    output_excel_path = f"{output_dir}/output_{unique_id}.xlsx"  # Save in the specified folder
    with pd.ExcelWriter(output_excel_path) as writer:
        processed_tables[0].to_excel(writer, sheet_name="Table_1", index=False)

        sheet_name = "Table_1"
        row_offset = processed_tables[0].shape[0] + 1

        for i in range(1, len(processed_tables)):
            current_table = processed_tables[i]
            if processed_tables[i-1].shape[1] == current_table.shape[1]:
                current_table.to_excel(writer, sheet_name=sheet_name, startrow=row_offset, index=False)
                row_offset += current_table.shape[0]
            elif processed_tables[i-1].shape[0] == current_table.shape[0]:
                current_table.to_excel(writer, sheet_name=sheet_name, startrow=0, startcol=processed_tables[i-1].shape[1], index=False)
            else:
                sheet_name = f"Table_{i+1}"
                current_table.to_excel(writer, sheet_name=sheet_name, index=False)
                row_offset = current_table.shape[0] + 1
    return output_excel_path


def merge_sheets_based_on_dimensions(excel_file_path, output_excel_path):
    """Merge sheets based on dimensions and save to the final Excel file"""
    excel_file = pd.ExcelFile(excel_file_path)
    sheet_names = excel_file.sheet_names

    row_offset = 0
    col_offset = 0

    with pd.ExcelWriter(output_excel_path) as writer:
        sheet_df = excel_file.parse(sheet_names[0])
        sheet_df = sheet_df.iloc[:-1]
        sheet_df.columns = sheet_df.columns.where(~sheet_df.columns.str.contains('^Unnamed'), '')
        sheet_df.to_excel(writer, sheet_name="Merged_Table", index=False, startrow=row_offset)
        row_offset += sheet_df.shape[0]

        for i in range(1, len(sheet_names)):
            sheet_df = excel_file.parse(sheet_names[i])
            sheet_df = sheet_df.iloc[:-1]
            sheet_df.columns = sheet_df.columns.where(~sheet_df.columns.str.contains('^Unnamed'), '')

            prev_sheet_df = excel_file.parse(sheet_names[i - 1])
            prev_sheet_df = prev_sheet_df.iloc[:-1]

            if prev_sheet_df.shape[0] == sheet_df.shape[0]:
                sheet_df.to_excel(writer, sheet_name="Merged_Table", index=False, startrow=0, startcol=prev_sheet_df.shape[1])
                col_offset = prev_sheet_df.shape[1] + sheet_df.shape[1]
            elif prev_sheet_df.shape[1] == sheet_df.shape[1]:
                sheet_df.to_excel(writer, sheet_name="Merged_Table", index=False, startrow=row_offset)
                row_offset += sheet_df.shape[0]
            else:
                sheet_df.to_excel(writer, sheet_name=f"Table_{i+1}", index=False)
                row_offset = sheet_df.shape[0] + 1


def main(input_pdf_path):

    output_dir = "attachments"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    intermediate_excel_path = process_pdf_to_excel(input_pdf_path, output_dir)
    if not intermediate_excel_path:
        print("Failed to process PDF to Excel.")
        return None

    unique_id = str(uuid.uuid4())
    final_excel_path = f"{output_dir}/final_{unique_id}.xlsx"  

    merge_sheets_based_on_dimensions(intermediate_excel_path, final_excel_path)

    return final_excel_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python script.py <input_pdf_path>")
        sys.exit(1)

    input_pdf_path = sys.argv[1]
    output_excel = main(input_pdf_path)
    if output_excel:
        print(output_excel)
